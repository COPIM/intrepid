import os
import tempfile
import zipfile
from decimal import Decimal
from typing import Any

import babel.numbers
from django.conf import settings
from django.contrib.auth.models import User
from django.db.models import Q
from django.shortcuts import reverse
from django.template.loader import render_to_string
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches

from accounts import models as account_models
from mail import models as mm
from package import models


def explain_price(
    package, identifier, identifier_type
) -> (
    tuple[int, None, str]
    | tuple[int, None]
    | tuple["models.Price", "models.Banding", str]
    | tuple["models.Price", "models.Banding"]
):
    """
    This function is used to explain the price calculation for a given package
    :param package: the package to calculate the price for
    :param identifier: the user or session to calculate the price for
    :param identifier_type: the type of identifier (user or session)
    :return: a tuple of the price, banding and explanation
    """
    country = identifier.profile.default_currency if identifier else None
    return get_price_for_package(
        package, identifier, identifier_type, country, explain=True
    )


def get_price_for_package(
    package, identifier, identifier_type, country=None, explain=False
) -> (
    tuple[int, None, str]
    | tuple[int, None]
    | tuple["models.Price", "models.Banding", str]
    | tuple["models.Price", "models.Banding"]
):
    """
    This function is used to calculate the price for a given package
    :param package: the package to calculate the price for
    :param identifier: the user or session to calculate the price for
    :param identifier_type: the type of identifier (user or session)
    :param country: the country to calculate the price for
    :param explain: whether to return an explanation of the price calculation
    :return: a tuple of the price, banding and explanation
    """

    # priority here should be:
    # 1. If the user has a preferred currency, find the first banding that
    # matches it
    # 2. Otherwise, fall back to the default currency/banding

    band_filter = None

    explanation = ""

    band_filter, bandings, country, explanation = get_band_filter(
        band_filter, country, explanation, package
    )

    banding, banding_choice = None, None

    fte = None

    banding_choice, explanation, fte = get_user(
        band_filter,
        banding_choice,
        explanation,
        fte,
        identifier,
        identifier_type,
        package,
    )

    if not fte:
        # if we get here, it's because it's a session basket and we don't
        # have any FTE info for the user
        explanation += "We don't have the data to calculate a price.\n"
        if explain:
            return 0, None, explanation
        else:
            return 0, None

    if band_filter and band_filter.banding_type.is_fte:
        # if we get here, the user has set a default currency, which we have
        # found and it's an FTE type
        banding, explanation = check_default_currency_and_fte(
            band_filter, bandings, explanation, fte, package
        )
    elif band_filter and not band_filter.banding_type.is_fte:
        # if we get here, the user has set a default currency, which we have
        # found and it's not an FTE type
        explanation += (
            "We have found a banded pricing in the customer's "
            "preferred currency.\n"
        )
        try:
            banding, explanation = user_without_banding(
                band_filter,
                banding,
                banding_choice,
                bandings,
                explanation,
                package,
            )
        except (ValueError, models.Banding.DoesNotExist):
            explanation += "Oh dear. We failed to retrieve a banding..\n"
            if explain:
                return 0, None, explanation
            else:
                return 0, None

    elif package.banding_type.is_fte:
        # this is falling back to the package default on an FTE calc
        banding, explanation = fallback_to_default(
            banding, bandings, explanation, fte, package
        )
    elif banding_choice:
        # this is falling back to the package default on a banded calc
        explanation += (
            "We are using the default banded pricing for this "
            "package because we couldn't find one in the user's "
            "preferred currency.\n"
        )
        try:
            banding = bandings.get(
                banding_type=package.banding_type,
                package=package,
                vocab=banding_choice,
            )
        except (ValueError, models.Banding.DoesNotExist):
            explanation += "It looks like pricing data is not set properly.\n"

            if explain:
                return 0, None, explanation
            else:
                return 0, None

    if banding:
        try:
            explanation, price = fetch_price(banding, country, explanation)
        except models.Price.DoesNotExist:
            # if we land here, it's because the currency requested by the
            # user does not exist. In this case, we use the default currency
            # for the package
            try:
                explanation, price = use_default_pricing(banding, explanation)
            except models.Price.DoesNotExist:
                # if we land here it's because the currency requested by the
                # user does not exist but there is also no pricing information
                # set for the default currency for the package.
                explanation += "No default pricing was set.\n"

                if explain:
                    return 0, None, explanation
                else:
                    return 0, None
        explanation += "Returning price {} for {}.\n".format(price, banding)

        if explain:
            return price, banding, explanation
        else:
            return price, banding
    else:
        explanation += (
            "We selected the default banding, but the user has "
            "not classified themselves so we can't show a price.\n"
        )

        if explain:
            return 0, None, explanation
        else:
            return 0, None


def fetch_price(banding, country, explanation) -> tuple[str, "models.Price"]:
    """
    This function is used to fetch a price
    :param banding: the banding to fetch the price for
    :param country: the country to fetch the price for
    :param explanation: the explanation of the price calculation
    :return: a tuple of the explanation and the price
    """
    explanation += "Trying to fetch {} for {}.\n".format(banding, country)
    price = models.Price.objects.get(
        banding=banding,
        country=country,
    )
    return explanation, price


def use_default_pricing(banding, explanation) -> tuple[str, "models.Price"]:
    """
    This function is used to use the default pricing
    :param banding: the banding to use the default pricing for
    :param explanation: the explanation of the price calculation
    :return: a tuple of the explanation and the price
    """
    explanation += "Trying to fetch {} for {} (default).\n".format(
        banding, banding.package.default_country
    )
    price = models.Price.objects.get(
        banding=banding,
        country=banding.package.default_country,
    )
    return explanation, price


def fallback_to_default(
    banding, bandings, explanation, fte, package
) -> tuple["models.Banding", str]:
    """
    This function is used to fall back to the default pricing
    :param banding: the banding to fall back to the default pricing for
    :param bandings: the bandings to fall back to the default pricing for
    :param explanation: the explanation of the price calculation
    :param fte: the FTE to fall back to the default pricing for
    :param package: the package to fall back to the default pricing for
    :return: a tuple of the banding and the explanation
    """
    explanation += (
        "We are using the default FTE pricing for this "
        "package because we couldn't find one in the user's "
        "preferred currency.\n"
    )
    banding = bandings.filter(
        Q(vocab__upper_limit=None) | Q(vocab__upper_limit__gte=fte),
        package=package,
        banding_type=package.banding_type,
        vocab__lower_limit__lte=fte,
    ).first()
    return banding, explanation


def get_user(
    band_filter,
    banding_choice,
    explanation,
    fte,
    identifier,
    identifier_type,
    package,
) -> tuple[str, str, Decimal]:
    """
    This function is used to get the user
    :param band_filter: the band filter
    :param banding_choice: the banding choice
    :param explanation: the explanation of the price calculation
    :param fte: the FTE
    :param identifier: the identifier
    :param identifier_type: the identifier type
    :param package: the package
    :return: a tuple of the banding choice, explanation and FTE
    """
    if identifier_type == "user":
        banding_choice, explanation, fte = find_user_in_database(
            band_filter, explanation, identifier, package
        )
    elif identifier_type == "session":
        banding_choice, explanation, fte = find_user_from_session(
            band_filter, explanation, identifier, package
        )
    return banding_choice, explanation, fte


def get_band_filter(
    band_filter, country, explanation, package
) -> tuple[
    "models.Banding", "models.QuerySet[models.Banding]", "models.Country", str
]:
    """
    This function is used to get the band filter
    :param band_filter: the band filter
    :param country: the country
    :param explanation: the explanation of the price calculation
    :param package: the package
    :return: a tuple of the band filter, country and explanation
    """
    band_filter, explanation = test_country(
        band_filter, country, explanation, package
    )
    if band_filter:
        bandings, explanation = handle_band_filter(
            band_filter, explanation, package
        )
    else:
        # If we get here, we need to check for a catch-all price
        band_filter, bandings, country, explanation = check_for_catch_all(
            band_filter, country, explanation, package
        )
    return band_filter, bandings, country, explanation


def user_without_banding(
    band_filter, banding, banding_choice, bandings, explanation, package
) -> tuple["models.Banding", str]:
    """
    This function is used to handle a user without a banding
    :param band_filter: the band filter
    :param banding: the banding object
    :param banding_choice: the banding choice
    :param bandings: the bandings
    :param explanation: the explanation of the price calculation
    :param package: the package
    :return: a tuple of the banding and the explanation
    """
    if banding_choice:
        banding = bandings.get(
            banding_type=band_filter.banding_type,
            package=package,
            vocab=banding_choice,
        )
    else:
        explanation += (
            "The user has not set their banding for this "
            "type of banding so we can't price. \n"
        )
        raise ValueError
    return banding, explanation


def check_default_currency_and_fte(
    band_filter, bandings, explanation, fte, package
) -> tuple["models.Banding", str]:
    """
    This function is used to check the default currency and FTE
    :param band_filter: the band filter
    :param bandings: the bandings
    :param explanation: the explanation of the price calculation
    :param fte: the FTE
    :param package: the package
    :return: a tuple of the banding and the explanation
    """
    explanation += (
        "We have found an FTE pricing in the customer's "
        "preferred currency.\n"
    )
    banding = bandings.filter(
        Q(vocab__upper_limit=None) | Q(vocab__upper_limit__gte=fte),
        package=package,
        banding_type=band_filter.banding_type,
        vocab__lower_limit__lte=fte,
    ).first()
    return banding, explanation


def check_for_catch_all(
    band_filter, country, explanation, package
) -> tuple[
    "models.Banding", "models.QuerySet[models.Banding]", "models.Country", str
]:
    """
    This function is used to check for a catch-all price
    :param band_filter: the band filter
    :param country: the country
    :param explanation: the explanation of the price calculation
    :param package: the package
    :return: a tuple of the band filter, bandings, country and explanation
    """
    if not band_filter:
        if country:
            # see if we can find a price that matches the desired code
            band_filter, country, explanation = match_price_to_code(
                band_filter, country, explanation, package
            )

        if band_filter:
            bandings, explanation = explain_catch_all_band_filter(
                band_filter, explanation, package
            )
        else:
            bandings = models.Banding.objects.filter(
                package=package,
                banding_type=package.banding_type,
            )
    else:
        bandings = models.Banding.objects.filter(
            package=package,
            banding_type=package.banding_type,
        )
    return band_filter, bandings, country, explanation


def find_user_from_session(
    band_filter, explanation, identifier, package
) -> tuple["models.Banding", str, Decimal]:
    """
    This function is used to find a user from a session
    :param band_filter: the band filter
    :param explanation: the explanation of the price calculation
    :param identifier: the identifier
    :param package: the package
    :return: a tuple of the banding choice, explanation and FTE
    """

    explanation += "We are looking up a user from the request session\n"
    fte = identifier.get("fte")
    if band_filter:
        session_string = "banding_type_{}".format(band_filter.banding_type.pk)
    else:
        session_string = "banding_type_{}".format(package.banding_type.pk)
    banding_choice = identifier.get(session_string, None)
    return banding_choice, explanation, fte


def find_user_in_database(
    band_filter, explanation, identifier, package
) -> tuple["models.Banding", str, Decimal]:
    """
    This function is used to find a user in the database
    :param band_filter: the band filter
    :param explanation: the explanation of the price calculation
    :param identifier: the identifier
    :param package: the package
    :return: a tuple of the banding choice, explanation and FTE
    """
    explanation += "We are looking up a user from the database\n"
    fte = identifier.profile.fte
    if band_filter:
        banding_choice = account_models.AccountBandingChoices.objects.filter(
            account=identifier,
            banding_type=band_filter.banding_type,
        ).first()
    else:
        banding_choice = account_models.AccountBandingChoices.objects.filter(
            account=identifier,
            banding_type=package.banding_type,
        ).first()
    if banding_choice:
        banding_choice = banding_choice.banding_type_vocab.pk
    return banding_choice, explanation, fte


def explain_catch_all_band_filter(
    band_filter, explanation, package
) -> tuple["models.Banding", str]:
    """
    This function is used to explain a catch-all band filter
    :param band_filter: the band filter
    :param explanation: the explanation of the price calculation
    :param package: the package
    :return: a tuple of the banding and the explanation
    """
    explanation += (
        "We found a catch-all band filter that "
        "can handle this: "
        "{}\n".format(band_filter)
    )
    bandings = models.Banding.objects.filter(
        package=package,
        banding_type=band_filter.banding_type,
    )
    return bandings, explanation


def match_price_to_code(
    band_filter, country, explanation, package
) -> tuple["models.Banding", "models.Country", str]:
    """
    This function is used to match a price to a code
    :param band_filter: the band filter
    :param country: the country (code)
    :param explanation: the explanation of the price calculation
    :param package: the package
    :return: a tuple of the band filter, country and explanation
    """
    explanation += "Attempting to find a catch-all price for " "{}\n".format(
        country
    )
    country_prices = package.price_bandings
    for key, val in country_prices.items():
        for inner_key, inner_val in val.items():
            if inner_key.catch_all:
                if inner_key.currency == country.currency:
                    explanation += "Using {}\n".format(inner_key)
                    band_filter = key.banding_type_entry
                    # switch over the country
                    country = inner_key
    return band_filter, country, explanation


def handle_band_filter(
    band_filter, explanation, package
) -> tuple["models.Banding", str]:
    """
    This function is used to handle a band filter
    :param band_filter: the band filter
    :param explanation: the explanation of the price calculation
    :param package: the package
    :return: a tuple of the banding and the explanation
    """
    explanation += (
        "We found a band filter that can handle this: "
        "{}\n".format(band_filter)
    )
    bandings = models.Banding.objects.filter(
        package=package,
        banding_type=band_filter.banding_type,
    )
    return bandings, explanation


def test_country(
    band_filter, country, explanation, package
) -> tuple["models.Banding", str]:
    """
    This function is used to test the country
    :param band_filter: the band filter
    :param country: the country
    :param explanation: the explanation of the price calculation
    :param package: the package
    :return: a tuple of the band filter and the explanation
    """
    if country:
        # see if we can find a price that matches the desired code
        explanation += (
            "The user in question prefers the currency " "{}\n".format(country)
        )
        country_prices = package.price_bandings

        for key, val in country_prices.items():
            for inner_key, inner_val in val.items():
                if inner_key == country:
                    band_filter = key.banding_type_entry
    return band_filter, explanation


def provisional_page_table_rows(order) -> dict[int, dict[str, Any]]:
    """
    This function is used to generate the table rows for the provisional page
    :param order: the order to generate the table rows for
    :return: a dictionary of table rows
    """
    rows = {}

    for signup in order.packagesignup_set.all():
        init_id = signup.associated_package.initiative.pk
        if not rows.get(init_id):
            rows[init_id] = {
                "initiative": signup.associated_package.initiative,
                "packages": [signup.associated_package],
                "contact_email": signup.associated_package.initiative.contact_email,
            }
        else:
            rows[init_id]["packages"].append(
                signup.associated_package,
            )

    return rows


def unlink_frozen_doc(order) -> None:
    """
    This function is used to unlink a frozen document
    :param order: the order to unlink the frozen document for
    :return: None
    """
    # if the given order has a document, that is a file on disk: unlink it.
    if order.document:
        file_path = order.document.final_zip.path
        if os.path.isfile(file_path):
            os.unlink(file_path)


def generate_package_docs_zip(
    order, request
) -> "models.FrozenPackageDocument":
    """
    This function is used to generate a zip file of the documents for a given
    :param order: the order to generate the zip file for
    :param request: the request object
    :return: the frozen document
    """
    # create temporary named file on the system
    with tempfile.NamedTemporaryFile(suffix=".zip") as f:
        # open a zip archive handle in write mode on the file
        with zipfile.ZipFile(f.name, "w") as zip_file:
            # get a list of documents

            # iterate over the document files in the package and add them
            # to the zip
            for package in order.package_set:
                try:
                    custom_package = models.CustomPackageDocument.objects.get(
                        package_signup__associated_package=package,
                        package_signup__associated_order=order,
                        pdf_file__isnull=False,
                    )
                    final_path = os.path.join(
                        package.initiative.name
                        + " ("
                        + str(package.initiative.pk)
                        + ")",
                        package.name + " (" + str(package.pk) + ")",
                        os.path.basename(custom_package.pdf_file.path),
                    )
                    zip_file.write(
                        filename=custom_package.pdf_file.path,
                        arcname=final_path,
                    )
                except (models.CustomPackageDocument.DoesNotExist, ValueError):
                    for document in package.packagedocument_set.all():
                        try:
                            if document.pdf_file:
                                with open(document.pdf_file.path, "rb"):
                                    # TODO: check that this context manager
                                    #  should be unused.

                                    # create a filename that merges the package
                                    # ID with the document filename, so that
                                    # packages can't overwrite each other's
                                    # documents
                                    final_path = os.path.join(
                                        package.initiative.name
                                        + " ("
                                        + str(package.initiative.pk)
                                        + ")",
                                        package.name
                                        + " ("
                                        + str(package.pk)
                                        + ")",
                                        os.path.basename(
                                            document.pdf_file.path
                                        ),
                                    )

                                    # write the file into the zip (filename is the file
                                    # and arcname is the name/path of the file in the
                                    # zip)
                                    zip_file.write(
                                        filename=document.pdf_file.path,
                                        arcname=final_path,
                                    )
                        except FileNotFoundError:
                            pass

        # now add the zip file to a frozen document
        frozen_document = models.FrozenPackageDocument()
        frozen_document.subject_line = "Documents for order {0} [{1}]".format(
            order.order_number,
            order.valid_period,
        )

        if request.user.is_authenticated:
            frozen_document.associated_user = request.user
        frozen_document.associated_session = request.session.session_key
        frozen_document.save()
        frozen_document.final_zip.save("{0}.zip".format(order.order_number), f)
        frozen_document.save()
        return frozen_document


def check_basket_for_disabled_signups(request, basket) -> str | None:
    """
    This function is used to check a basket for disabled signups
    :param request: the request object
    :param basket: the basket to check
    :return: an error message if there is one, otherwise None
    """
    if not request.site.enable_signup:
        return "Signup is disabled for all packages across the site."
    elif (
        not request.site.enable_meta_package_signup
        and basket.meta_packages.all()
    ):
        return (
            "Collective signup is currently disabled and you have a "
            "collective package in your quote."
        )
    elif (
        not request.site.enable_individual_package_signup
        and basket.packages.all()
    ):
        return (
            "Individual initiative signup is currently disabled and you "
            "have an initiative package in your quote."
        )

    return None


def send_new_order_notification(
    order, request, identifier, identifier_type
) -> None:
    """
    This function is used to send a new order notification
    :param order: the order to send the notification for
    :param request: the request object
    :param identifier: the identifier
    :param identifier_type: the identifier type
    :return: None
    """
    # Notify OBC
    email_template = mm.EmailTemplate.objects.get(
        name="platform_new_order",
    )
    url = request.build_absolute_uri(
        reverse(
            "order_detail",
            kwargs={
                "order_id": order.pk,
            },
        )
    )
    context = {
        "url": url,
        "order": order,
        "request": request,
    }
    email_template.send(
        to=request.site.contact_email,
        context=context,
    )

    # Notify orderer

    quote_path, quote_filename = generate_doc(
        request,
        order,
        identifier,
        identifier_type,
    )
    accept_path, accept_filename = generate_acceptance_doc(
        request,
        order,
        identifier,
        identifier_type,
    )

    email_template = mm.EmailTemplate.objects.get(
        name="order_complete",
    )
    url = request.build_absolute_uri(
        reverse(
            "order_provisional",
            kwargs={
                "order_id": order.pk,
            },
        )
    )
    context = {
        "request": request,
        "url": url,
        "order": order,
    }
    email_template.send(
        to=(
            order.associated_user.email
            if order.associated_user
            else order.email_address
        ),
        context=context,
        attachments=[quote_path, accept_path],
    )


def send_order_complete_notification(order, request) -> None:
    """
    This function is used to send an order complete notification
    :param order: the order to send the notification for
    :param request: the request object
    :return: None
    """
    recipients = []
    email = mm.EmailTemplate.objects.get(
        name="order_complete",
    )

    # Build list of initiative reps, billing managers and order owners:
    for signup in order.packagesignup_set.all():
        initiative_url = request.build_absolute_uri(
            reverse(
                "process_signup",
                kwargs={
                    "package_id": signup.associated_package.pk,
                    "initiative_id": signup.associated_package.initiative.pk,
                    "signup_id": signup.pk,
                },
            )
        )

        for user in signup.associated_package.initiative.users.all():
            recipients.append(
                {
                    "to": user.email,
                    "recipient": user,
                    "type": "initiative",
                    "url": initiative_url,
                }
            )

        try:
            banding_type_currency_entry = models.BandingTypeCurrencyEntry.objects.get(
                package=signup.associated_package,
                country=signup.currency,
                banding_type_entry__banding_type=signup.banding.banding_type,
            )

            payment_processor = banding_type_currency_entry.payment_processor
            for user in payment_processor.managers.all():
                payment_processor_url = request.build_absolute_uri(
                    reverse(
                        "detail_invoice",
                        kwargs={
                            "payment_processor_id": payment_processor.pk,
                            "order_id": order.pk,
                        },
                    )
                )
                recipients.append(
                    {
                        "to": user.email,
                        "recipient": user,
                        "type": "billing_manager",
                        "url": payment_processor_url,
                    }
                )
        except (AttributeError, models.BandingTypeCurrencyEntry.DoesNotExist):
            print("Nothing found here.")

    recipients.append(
        {
            "to": order.associated_user.email,
            "recipient": order.associated_user,
            "type": "library",
            "url": request.build_absolute_uri(
                reverse(
                    "order_complete",
                    kwargs={
                        "order_id": order.pk,
                    },
                )
            ),
        }
    )

    for user in User.objects.filter(
        is_staff=True,
    ):
        recipients.append(
            {
                "to": user.email,
                "recipient": user,
                "type": "staff",
                "url": request.build_absolute_uri(
                    reverse(
                        "order_detail",
                        kwargs={
                            "order_id": order.pk,
                        },
                    )
                ),
            }
        )

    for recipient in recipients:
        context = {
            "url": recipient.get("url"),
            "request": request,
            "type": recipient.get("type"),
            "recipient": recipient.get("recipient"),
            "order": order,
        }
        email.send(
            to=recipient.get("to"),
            context=context,
        )


def format_price(cost, currency) -> str:
    """
    This function is used to format a price
    :param cost: the cost to format
    :param currency: the currency to format
    :return: the formatted price
    """
    try:
        return babel.numbers.format_currency(cost, currency, locale="en_US")
    except (
        babel.numbers.UnknownCurrencyError,
        babel.numbers.UnknownCurrencyFormatError,
        babel.numbers.NumberFormatError,
        ValueError,
    ):
        return "{} {}".format(cost, currency)


def calculate_site_percentage(
    request, total
) -> tuple[dict[str, Decimal], Decimal] | tuple[dict[str, Decimal], None]:
    """
    This function is used to calculate the site percentage
    :param request: the request object
    :param total: the total to calculate the site percentage for
    :return: a tuple of the site percentage and the site percentage value
    """

    """Assumes a single key value pair."""
    for currency, total in total.items():
        site_percentage = Decimal(
            round((float(request.site.fee_amount) / 100) * float(total), 2)
        )
        return {
            currency: site_percentage,
        }, site_percentage

    return {}, None


def filter_duplicates_from_costs(costs) -> list[dict[str, Any]]:
    """
    This function is used to filter duplicates from costs
    :param costs: the costs to filter
    :return: the filtered costs
    """
    costs_to_return = []
    checked_packages = []

    for cost in costs:
        if cost.get("package") not in checked_packages:
            costs_to_return.append(cost)
            checked_packages.append(cost.get("package"))

    return costs_to_return


def convert_currency_totals(
    request, identifier_type, identifier, totals
) -> (
    tuple[None, dict, None, None, None]
    | tuple[str, dict, str, Decimal, Decimal]
):
    if identifier_type == "user":
        user = account_models.User.objects.get(username=identifier)
        country = user.profile.default_currency
    else:
        country_pk = identifier.get("currency", None)
        country = models.Country.objects.filter(pk=country_pk).first()

    from package import currency as convert_currency

    if country:
        amounts = []
        for currency, value in totals.items():
            if currency == country.currency:
                amounts.append(value)
            else:
                try:
                    converted_currency = convert_currency.convert(
                        currency_from=currency,
                        currency_to=country.currency,
                        value=value,
                    )
                    amounts.append(converted_currency)
                except ValueError:
                    # the country selected does not have a conversion rate.
                    country = models.Country.objects.get(
                        currency=request.site.fallback_currency,
                    )
                    converted_currency = convert_currency.convert(
                        currency_from=currency,
                        currency_to=country.currency,
                        value=value,
                    )
                    amounts.append(converted_currency)

        currency_conversion_total = Decimal(round(sum(amounts), 2))
        site_percentage, site_percentage_value = calculate_site_percentage(
            request, {country.currency: currency_conversion_total}
        )

        return (
            site_percentage,
            {country.currency: currency_conversion_total},
            country.currency,
            currency_conversion_total,
            site_percentage_value,
        )
    return None, {}, None, None, None


def generate_doc(
    request, order, identifier, identifier_type
) -> tuple[str, str]:
    """
    This function is used to generate a document
    :param request: the request object
    :param order: the order to generate the document for
    :param identifier: the identifier
    :param identifier_type: the identifier type
    :return: a tuple of the path and filename of the document
    """

    total = ""
    processing_fee = ""

    converted_currency = None
    package_costs, currency_totals = order.basket.cost(
        identifier=identifier, identifier_type=identifier_type
    )

    if len(currency_totals) > 1:
        (
            site_percentage,
            converted_total,
            converted_currency,
            currency_conversion_total,
            site_percentage_value,
        ) = convert_currency_totals(
            request=request,
            identifier_type=identifier_type,
            identifier=identifier,
            totals=currency_totals,
        )

    if order.converted_currency and order.converted_value:
        processing_fee_calc = order.platform_fee * order.term

        processing_fee = format_price(
            processing_fee_calc, order.converted_currency
        )
        total = "{}".format(
            format_price(
                (
                    order.converted_value * order.term
                    + order.platform_fee * order.term
                ),
                order.converted_currency,
            )
        )
    else:
        costs, currency_totals = order.basket.cost(identifier, identifier_type)
        for k, v in currency_totals.items():
            processing_fee = format_price(
                int(order.platform_fee) * int(order.term), k
            )
            total = "{}".format(
                format_price(
                    (
                        int(v) * int(order.term)
                        + int(order.platform_fee) * int(order.term)
                    ),
                    k,
                )
            )
    css = [
        os.path.join(
            settings.BASE_DIR, "static", "frontend", "css", "bootstrap.min.css"
        ),
    ]
    context = {
        "order": order,
        "total": total,
        "processing_fee": processing_fee,
        "css": css,
        "package_costs": package_costs,
        "converted_currency": converted_currency,
    }

    rendered_html = render_to_string("docs/order_doc.html", context)
    import pdfkit

    options = {
        "enable-local-file-access": None,
    }
    file_name = f"OBC_Quote_{order.pk}.pdf"
    save_path = os.path.join(settings.BASE_DIR, "files", "private_documents")
    full_path = os.path.join(save_path, file_name)
    pdfkit.from_string(rendered_html, full_path, css=css, options=options)

    return full_path, file_name


def generate_acceptance_doc(
    request, order, identifier, identifier_type
) -> tuple[str, str]:
    """
    This function is used to generate an acceptance document
    :param request: the request object
    :param order: the order to generate the acceptance document for
    :param identifier: the identifier
    :param identifier_type: the identifier type
    :return: a tuple of the path and filename of the acceptance document
    """

    # TODO: refactor long function

    order_doc = Document()
    order_doc.add_picture(
        os.path.join(
            settings.BASE_DIR, "static", "frontend", "img", "obc_logo.png"
        ),
        width=Inches(2.5),
    )
    order_doc.add_heading("Acceptance Form: Summary", 0)

    p = order_doc.add_paragraph()
    quote_runner = p.add_run(f"Quotation Reference Number: {order.pk}")
    quote_runner.bold = True

    for detail in order.get_order_form_details():
        order_doc.add_paragraph(
            "{}: {}".format(
                detail.get("question"),
                detail.get("answer"),
            )
        )
    order_doc.add_paragraph(f"Contact Email: {order.email_address}")

    order_doc.add_paragraph(
        "Please note, we may pass your name and email address to the Provider(s) you have chosen to support. If you do not want us to pass you on your details please make a note next to your contact name."
    )

    order_doc.add_paragraph(
        "You have selected the following offers to potentially support: "
    )

    for package in order.package_set:
        order_doc.add_paragraph(package.name, style="List Bullet")

    order_doc.add_paragraph(
        "As detailed in the Quotation document, the cost of your annual subscription is as follows: "
    )

    if order.converted_currency and order.converted_value:
        order_doc.add_paragraph(
            "Offers: {} {}".format(
                order.converted_currency, order.converted_value
            )
        )
        order_doc.add_paragraph(
            "OBC Processing Fee: {} {}".format(
                order.converted_currency, order.platform_fee
            )
        )
        total = "Total: {} {}".format(
            order.converted_currency,
            (order.converted_value + order.platform_fee),
        )
        add_bold_paragraph(
            order_doc,
            total,
        )
    else:
        package_costs, currency_totals = order.basket.cost(
            identifier=identifier, identifier_type=identifier_type
        )

        if len(currency_totals) > 1:
            (
                site_percentage,
                currency_totals,
                converted_currency,
                currency_conversion_total,
                site_percentage_value,
            ) = convert_currency_totals(
                request=request,
                identifier_type=identifier_type,
                identifier=identifier,
                totals=currency_totals,
            )

        for k, v in currency_totals.items():
            order_doc.add_paragraph(f"Offers: {k} {v}")
            order_doc.add_paragraph(
                f"OBC Processing Fee: {k} {order.platform_fee}"
            )
            total = "Total: {} {}".format(
                k,
                (v + order.platform_fee),
            )
            add_bold_paragraph(
                order_doc,
                total,
            )

    break_p = order_doc.add_paragraph()
    break_run = break_p.add_run()
    break_run.add_break(WD_BREAK.PAGE)

    order_doc.add_heading("Acceptance Form: Details", 0)
    order_doc.add_paragraph(
        "If you would like to proceed to support the chosen initiatives on the basis of the Quotation provided to you by the OBC, please add the details requested below to this document, editing as appropriate."
    )
    order_doc.add_paragraph(
        "Return the completed document to info@openbookcollective.org."
    )
    order_doc.add_paragraph(
        "Once we have reviewed the additional details, and if no further information is required, we will go ahead and issue an invoice."
    )
    add_bold_paragraph(
        order_doc,
        "Additional Contact Name(s): ",
        "[optional – enter as required]",
    )
    add_bold_paragraph(
        order_doc,
        "Additional Contact Email(s): ",
        "[optional – enter as required]",
    )
    add_bold_paragraph(
        order_doc,
        "Support Period (delete as appropriate): ",
        "1 year / 2 years / 3 years",
    )
    add_bold_paragraph(
        order_doc,
        "Start date of Support: ",
        "[insert date]",
    )
    add_bold_paragraph(
        order_doc,
        "In order to recognise the support it receives, the OBC would, where possible, like to list Supporting Institutions on its website and publicise this support via other channels (e.g. in newsletters, on social media). Do you give permission for the OBC to acknowledge your Institution’s support in this way?",
    )
    order_doc.add_paragraph("Yes / No")
    add_bold_paragraph(
        order_doc,
        "Do you give permission for the OBC and/or supported initiatives to use your institution's logo to acknowledge your institution’s support on their channels -- including on respective websites and social media channels, as well as in presentations? ",
    )
    add_bold_paragraph(
        order_doc,
        "Once we have reviewed this form, would you like the OBC to invoice you immediately, or should we contact you to confirm further invoicing details (e.g. obtain a purchase order)? (Delete as appropriate)",
    )
    order_doc.add_paragraph(
        "Invoice immediately / Obtain further invoicing details "
    )
    add_bold_paragraph(
        order_doc,
        "On behalf of the Institution listed in the above Summary, and from the Start Date and for the Support Period stated in this Acceptance Form, I accept the annual cost for supporting the chosen Offers as listed in the Quotation document (reference number detailed above), as well as the OBC terms and conditions: ",
    )
    order_doc.add_paragraph("Yes / No ")
    add_bold_paragraph(
        order_doc,
        "Name: ",
    )
    add_bold_paragraph(
        order_doc,
        "Date: ",
    )
    order_doc.add_paragraph(
        "Return this form to us at info@openbookcollective.org. You can also use this email address to contact us with any questions you have.",
    )

    filename = "OBC_Acceptance_Form_{}.docx".format(order.pk)
    save_path = os.path.join(
        settings.BASE_DIR, "files", "private_documents", filename
    )
    order_doc.save(save_path)

    return save_path, filename


def add_bold_paragraph(document, text, secondary_text=None) -> None:
    """
    This function is used to add a bold paragraph to a document
    :param document: the document to add the paragraph to
    :param text: the text to add
    :param secondary_text: the secondary text to add
    :return: None
    """
    p = document.add_paragraph()
    runner = p.add_run(text)
    runner.bold = True
    if secondary_text:
        p.add_run(secondary_text)


def get_user_currency(identifier, identifier_type) -> "models.Country":
    """
    This function is used to get a user's currency
    :param identifier: the identifier
    :param identifier_type: the identifier type
    :return: the user's currency
    """
    if identifier_type == "user":
        user = account_models.User.objects.get(username=identifier)
        country = user.profile.default_currency
    else:
        country_pk = identifier.get("currency", None)
        country = models.Country.objects.filter(pk=country_pk).first()

    return country


def add_pre_calc_to_objects(country_code, packages):
    country = models.Country.objects.get(code=country_code)

    for package in packages:
        try:
            package.pre_calc = models.PreCalcMinMax.objects.get(
                country=country,
                package=package,
            )
        except models.PreCalcMinMax.DoesNotExist:
            try:
                catch_all_country = models.Country.objects.get(
                    catch_all=True,
                    currency=country.currency,
                )
                package.pre_calc = models.PreCalcMinMax.objects.get(
                    country=catch_all_country,
                    package=package,
                )
            except (
                models.Country.DoesNotExist,
                models.PreCalcMinMax.DoesNotExist,
            ):
                package.pre_calc = models.PreCalcMinMax.objects.get(
                    country=package.default_country,
                    package=package,
                )
    return packages


def add_pre_calc_to_meta_objects(country_code, packages):
    country = models.Country.objects.get(code=country_code)

    for package in packages:
        try:
            package.pre_calc = models.PreCalcMinMax.objects.get(
                country=country,
                meta_package=package,
            )
        except models.PreCalcMinMax.DoesNotExist:
            try:
                catch_all_country = models.Country.objects.get(
                    catch_all=True,
                    currency=country.currency,
                )
                package.pre_calc = models.PreCalcMinMax.objects.get(
                    country=catch_all_country,
                    meta_package=package,
                )
            except (
                models.Country.DoesNotExist,
                models.PreCalcMinMax.DoesNotExist,
            ):
                package.pre_calc = models.PreCalcMinMax.objects.get(
                    country=package.default_country,
                    meta_package=package,
                )
    return packages
